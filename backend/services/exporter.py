"""
NyayaLens Premium Legal Audit Report Generator
Produces a structured, professional DOCX with cover page, executive summary,
conflict analysis, and per-clause deep-dive including SHAP highlights,
regulations, case law precedents, and AI rewrites.
"""

import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Design Tokens ──────────────────────────────────────────────────────────────
CLR_PRIMARY  = RGBColor(0x31, 0x2E, 0x81)   # Indigo 900
CLR_ACCENT   = RGBColor(0x4F, 0x46, 0xE5)   # Indigo 600
CLR_DARK     = RGBColor(0x1F, 0x29, 0x37)   # Gray 800
CLR_MUTED    = RGBColor(0x6B, 0x72, 0x80)   # Gray 500
CLR_LIGHT    = RGBColor(0xF9, 0xFA, 0xFB)   # Gray 50
CLR_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
CLR_BORDER   = RGBColor(0xE5, 0xE7, 0xEB)   # Gray 200

RISK_CLR = {
    "high":   RGBColor(0xDC, 0x26, 0x26),
    "medium": RGBColor(0xD9, 0x77, 0x06),
    "low":    RGBColor(0x25, 0x63, 0xEB),
    "safe":   RGBColor(0x05, 0x96, 0x69),
}
RISK_BG = {
    "high":   "FEF2F2",
    "medium": "FFFBEB",
    "low":    "EFF6FF",
    "safe":   "F0FDF4",
}
RISK_LABEL = {
    "high": "HIGH RISK", "medium": "CAUTION", "low": "LOW RISK", "safe": "SAFE"
}
SEV_CLR = {
    "critical": RGBColor(0xDC, 0x26, 0x26),
    "warning":  RGBColor(0xD9, 0x77, 0x06),
    "info":     RGBColor(0x25, 0x63, 0xEB),
}
SEV_BG = {"critical": "FEF2F2", "warning": "FFFBEB", "info": "EFF6FF"}

# ── XML Helpers ────────────────────────────────────────────────────────────────

def _shade_cell(cell, hex_color: str):
    tc = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc.append(shd)

def _cell_pad(cell, t=100, b=100, l=150, r=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", t), ("bottom", b), ("left", l), ("right", r)]:
        n = OxmlElement(f"w:{side}")
        n.set(qn("w:w"), str(val))
        n.set(qn("w:type"), "dxa")
        tcMar.append(n)
    tcPr.append(tcMar)

def _cell_valign(cell, align="center"):
    tc = cell._tc.get_or_add_tcPr()
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), align)
    tc.append(v)

def _no_border_tbl(tbl):
    """Remove all borders from a table."""
    tbl_el = tbl._tbl
    # CT_Tbl has no get_or_add_tblPr — find or create manually
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tbl_pr.append(borders)

def _hr(doc, color="E5E7EB", before=Pt(10), after=Pt(10), size="4"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = before
    p.paragraph_format.space_after = after
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), size)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def _spacer(doc, size=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = size
    p.paragraph_format.space_after = Pt(0)

def _section_heading(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = CLR_PRIMARY
    r.font.name = "Arial"
    _hr(doc, "312E81", before=Pt(2), after=Pt(12), size="6")

def _add_footer(doc, doc_id: str):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(f"NyayaLens Legal Intelligence  ·  Ref: {doc_id}  ·  Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = CLR_MUTED

    # Build PAGE field correctly: fldChar(begin) → instrText → fldChar(end)
    def _fld_run(ftype, text=None):
        r = p.add_run()
        r.font.size = Pt(8)
        r.font.color.rgb = CLR_MUTED
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), ftype)
        r._r.append(fc)
        if text:
            instr = OxmlElement("w:instrText")
            instr.text = text
            r._r.append(instr)
        return r

    _fld_run("begin")
    _fld_run("separate", " PAGE ")
    _fld_run("end")

    conf = p.add_run("  ·  Confidential")
    conf.font.size = Pt(8)
    conf.font.color.rgb = CLR_MUTED

# ── Section Builders ───────────────────────────────────────────────────────────

def _build_cover(doc, data: dict):
    """Full-width cover page with branding, document name, and risk verdict."""
    # Brand header table
    tbl = doc.add_table(rows=1, cols=2)
    _no_border_tbl(tbl)
    tbl.autofit = False
    tbl.columns[0].width = Inches(4.0)
    tbl.columns[1].width = Inches(2.5)

    lc = tbl.rows[0].cells[0]
    _cell_pad(lc, 0, 0, 0, 0)
    lp = lc.paragraphs[0]
    brand = lp.add_run("NYAYALENS")
    brand.bold = True
    brand.font.size = Pt(28)
    brand.font.color.rgb = CLR_PRIMARY
    brand.font.name = "Arial"
    sub = lp.add_run("\nAI-Powered Legal Risk Intelligence")
    sub.font.size = Pt(10)
    sub.font.color.rgb = CLR_MUTED

    rc = tbl.rows[0].cells[1]
    _cell_pad(rc, 0, 0, 0, 0)
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ts = datetime.now().strftime("%d %B %Y\n%I:%M %p IST")
    tr = rp.add_run(f"GENERATED\n{ts}")
    tr.font.size = Pt(9)
    tr.font.color.rgb = CLR_MUTED

    _hr(doc, "312E81", before=Pt(16), after=Pt(20), size="8")

    # Document title
    file_name = str(data.get("fileName", "Contract Document"))
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(4)
    tr2 = title_p.add_run(file_name)
    tr2.bold = True
    tr2.font.size = Pt(20)
    tr2.font.color.rgb = CLR_DARK

    lang = data.get("language") or {}
    lang_name = lang.get("name", "English") if isinstance(lang, dict) else "English"
    version = data.get("version", 1)
    rewrites = data.get("rewritesApplied", 0)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(0)
    mr = meta_p.add_run(
        f"Language: {lang_name}  ·  Version: v{version}"
        + (f"  ·  {rewrites} AI Rewrite(s) Applied" if rewrites else "")
        + f"  ·  Ref: {data.get('documentId', 'N/A')}"
    )
    mr.font.size = Pt(9)
    mr.font.color.rgb = CLR_MUTED

    _spacer(doc, Pt(16))

    # Risk verdict + breakdown side by side
    risk_score = int(data.get("overallRiskScore", 0))
    verdict_key = "high" if risk_score >= 70 else ("medium" if risk_score >= 40 else "safe")
    breakdown = data.get("riskBreakdown") or {}

    vtbl = doc.add_table(rows=1, cols=5)
    vtbl.autofit = False  # Must be set BEFORE column widths
    _no_border_tbl(vtbl)
    vtbl.columns[0].width = Inches(1.8)
    for i in range(1, 5):
        vtbl.columns[i].width = Inches(1.175)

    # Score card
    sc = vtbl.rows[0].cells[0]
    _shade_cell(sc, RISK_BG[verdict_key])
    _cell_pad(sc, 180, 180, 160, 160)
    _cell_valign(sc, "center")
    sp = sc.paragraphs[0]
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lbl = sp.add_run("OVERALL RISK\n")
    lbl.font.size = Pt(8)
    lbl.font.color.rgb = CLR_MUTED
    sv = sp.add_run(f"{risk_score}")
    sv.bold = True
    sv.font.size = Pt(36)
    sv.font.color.rgb = RISK_CLR[verdict_key]
    sp.add_run("\n/100").font.size = Pt(9)
    vv = sp.add_run(f"\n{RISK_LABEL[verdict_key]}")
    vv.bold = True
    vv.font.size = Pt(9)
    vv.font.color.rgb = RISK_CLR[verdict_key]

    # Per-level breakdown cells
    for i, lvl in enumerate(["high", "medium", "low", "safe"]):
        count = breakdown.get(lvl, 0)
        cell = vtbl.rows[0].cells[i + 1]
        _shade_cell(cell, RISK_BG[lvl])
        _cell_pad(cell, 150, 150, 120, 120)
        _cell_valign(cell, "center")
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cv = cp.add_run(str(count))
        cv.bold = True
        cv.font.size = Pt(24)
        cv.font.color.rgb = RISK_CLR[lvl]
        cl = cp.add_run(f"\n{RISK_LABEL[lvl]}")
        cl.font.size = Pt(8)
        cl.font.color.rgb = RISK_CLR[lvl]
        cp.add_run("\nclauses").font.size = Pt(8)

    _spacer(doc, Pt(20))


def _build_summary(doc, data: dict):
    """AI-generated executive summary + key stats."""
    _section_heading(doc, "Executive Summary")

    summary = str(data.get("summary", "No summary available."))
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(12)
    sr = sp.add_run(summary)
    sr.font.size = Pt(10.5)
    sr.font.color.rgb = CLR_DARK

    # Key stats row
    clauses = data.get("clauses") or []
    conflicts = data.get("conflicts") or []
    proc_time = float(data.get("processingTime") or 0)  # None-safe
    regs = sum(len(c.get("regulations") or []) for c in clauses)
    precs = sum(len(c.get("retrievedPrecedents") or []) for c in clauses)

    stats = [
        ("Clauses Analysed", str(len(clauses))),
        ("Conflicts Found", str(len(conflicts))),
        ("Regulations Matched", str(regs)),
        ("Precedents Retrieved", str(precs)),
        ("Processing Time", f"{proc_time:.1f}s"),
    ]

    stbl = doc.add_table(rows=1, cols=len(stats))
    _no_border_tbl(stbl)
    stbl.autofit = False
    col_w = Inches(6.5 / len(stats))
    for i, col in enumerate(stbl.columns):
        col.width = col_w

    for i, (label, value) in enumerate(stats):
        cell = stbl.rows[0].cells[i]
        _shade_cell(cell, "F3F4F6")
        _cell_pad(cell, 120, 120, 120, 120)
        _cell_valign(cell, "center")
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = cp.add_run(value)
        vr.bold = True
        vr.font.size = Pt(18)
        vr.font.color.rgb = CLR_PRIMARY
        lr = cp.add_run(f"\n{label}")
        lr.font.size = Pt(8)
        lr.font.color.rgb = CLR_MUTED


def _build_conflicts(doc, data: dict):
    """Conflict pairs table with severity, description, and clause references."""
    conflicts = data.get("conflicts") or []
    if not conflicts:
        return

    _section_heading(doc, "Conflict Analysis")

    # Summary line
    n_crit = sum(1 for c in conflicts if c.get("severity") == "critical")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(10)
    sr = sp.add_run(
        f"{len(conflicts)} internal contradiction(s) detected — "
        f"{n_crit} critical. Do not execute this contract until all critical "
        f"conflicts are resolved."
    )
    sr.font.size = Pt(10)
    sr.font.color.rgb = CLR_DARK if n_crit == 0 else RISK_CLR["high"]
    sr.italic = (n_crit > 0)

    for conflict in conflicts:
        sev = conflict.get("severity", "info")
        sev_clr = SEV_CLR.get(sev, CLR_MUTED)
        sev_bg = SEV_BG.get(sev, "EFF6FF")

        ctbl = doc.add_table(rows=1, cols=1)
        _no_border_tbl(ctbl)
        cell = ctbl.rows[0].cells[0]
        _shade_cell(cell, sev_bg)
        _cell_pad(cell, 140, 140, 180, 180)

        # Header line: severity badge + clause titles
        hp = cell.paragraphs[0]
        hp.paragraph_format.space_after = Pt(4)
        badge = hp.add_run(f"[{sev.upper()}]  ")
        badge.bold = True
        badge.font.size = Pt(9)
        badge.font.color.rgb = sev_clr
        title = hp.add_run(
            f"\"{conflict.get('clauseATitle','')}\"  ⟷  \"{conflict.get('clauseBTitle','')}\" "
        )
        title.bold = True
        title.font.size = Pt(10)
        title.font.color.rgb = CLR_DARK

        # Description — added inside the cell, not to doc directly
        dp = cell.add_paragraph()
        dp.paragraph_format.space_before = Pt(2)
        dp.paragraph_format.space_after = Pt(0)
        dr = dp.add_run(str(conflict.get("description", "")))
        dr.font.size = Pt(9.5)
        dr.font.color.rgb = CLR_DARK

        # Spacer paragraph after the conflict table (not inside cell)
        _spacer(doc, Pt(8))


def _build_clauses(doc, data: dict):
    """Per-clause deep-dive: score, category, text, SHAP, regulations, precedents, rewrite."""
    clauses = data.get("clauses") or []
    if not clauses:
        return

    _section_heading(doc, "Clause-by-Clause Analysis")

    # Sort high-risk first for impact
    sorted_clauses = sorted(clauses, key=lambda c: -int(c.get("riskScore", 0)))

    for idx, clause in enumerate(sorted_clauses):
        lvl  = clause.get("riskLevel", "safe")
        clr  = RISK_CLR.get(lvl, CLR_DARK)
        bg   = RISK_BG.get(lvl, "FFFFFF")
        score = int(clause.get("riskScore", 0))
        title = str(clause.get("title", f"Clause {idx+1}"))
        category = str(clause.get("category", "General"))
        text  = str(clause.get("text", "[No text]"))

        # ── Title bar ──────────────────────────────────────────────────────
        tbar = doc.add_table(rows=1, cols=3)
        _no_border_tbl(tbar)
        tbar.autofit = False
        tbar.columns[0].width = Inches(0.6)
        tbar.columns[1].width = Inches(4.9)
        tbar.columns[2].width = Inches(1.0)

        # Left: coloured risk badge cell
        lc = tbar.rows[0].cells[0]
        _shade_cell(lc, bg)
        _cell_pad(lc, 100, 100, 120, 80)
        _cell_valign(lc, "center")
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(f"{score}")
        lr.bold = True
        lr.font.size = Pt(14)
        lr.font.color.rgb = clr

        # Middle: title + category
        mc = tbar.rows[0].cells[1]
        _shade_cell(mc, bg)
        _cell_pad(mc, 100, 100, 80, 80)
        _cell_valign(mc, "center")
        mp = mc.paragraphs[0]
        mr = mp.add_run(f"§{clause.get('index', idx)+1}  {title}")
        mr.bold = True
        mr.font.size = Pt(11)
        mr.font.color.rgb = CLR_DARK
        cat_r = mp.add_run(f"  ·  {category}")
        cat_r.font.size = Pt(9)
        cat_r.font.color.rgb = CLR_MUTED

        # Right: risk level label
        rc = tbar.rows[0].cells[2]
        _shade_cell(rc, bg)
        _cell_pad(rc, 100, 100, 80, 120)
        _cell_valign(rc, "center")
        rp = rc.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = rp.add_run(RISK_LABEL.get(lvl, lvl.upper()))
        rr.bold = True
        rr.font.size = Pt(9)
        rr.font.color.rgb = clr

        # ── Clause text ────────────────────────────────────────────────────
        tp = doc.add_paragraph()
        tp.paragraph_format.left_indent = Inches(0.15)
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.space_after = Pt(6)
        tr_ = tp.add_run(text)
        tr_.font.size = Pt(9.5)
        tr_.font.color.rgb = CLR_DARK

        # ── SHAP highlights ────────────────────────────────────────────────
        shap_tokens = clause.get("shapTokens") or []
        # Use .get() consistently to avoid KeyError on either dict or object
        top_risky = sorted(
            [t for t in shap_tokens if (t.get("score") or 0) > 0.3],
            key=lambda t: -(t.get("score") or 0)
        )[:6]
        if top_risky:
            shp = doc.add_paragraph()
            shp.paragraph_format.left_indent = Inches(0.15)
            shp.paragraph_format.space_before = Pt(2)
            shp.paragraph_format.space_after = Pt(4)
            head_run = shp.add_run("Risk Triggers: ")
            head_run.font.size = Pt(8.5)
            head_run.bold = True
            head_run.font.color.rgb = CLR_MUTED
            for tok in top_risky:
                word = tok.get("word") or ""
                score = tok.get("score") or 0
                wr = shp.add_run(f"  {word} (+{score:.2f})  ")
                wr.font.size = Pt(8.5)
                wr.font.color.rgb = RISK_CLR["high"]

        # ── Regulations ────────────────────────────────────────────────────
        regs = clause.get("regulations") or []
        if regs:
            rp2 = doc.add_paragraph()
            rp2.paragraph_format.left_indent = Inches(0.15)
            rp2.paragraph_format.space_before = Pt(2)
            rp2.paragraph_format.space_after = Pt(2)
            rhead = rp2.add_run("Applicable Law: ")
            rhead.font.size = Pt(8.5)
            rhead.bold = True
            rhead.font.color.rgb = CLR_MUTED
            for reg in regs[:4]:
                rr2 = rp2.add_run(
                    f"{reg.get('body','')} § {reg.get('code','')} — {reg.get('title','')}  "
                )
                rr2.font.size = Pt(8.5)
                rr2.font.color.rgb = CLR_ACCENT

        # ── Precedents ─────────────────────────────────────────────────────
        precs = clause.get("retrievedPrecedents") or []
        if precs:
            pp2 = doc.add_paragraph()
            pp2.paragraph_format.left_indent = Inches(0.15)
            pp2.paragraph_format.space_before = Pt(2)
            pp2.paragraph_format.space_after = Pt(4)
            phead = pp2.add_run("Case Law: ")
            phead.font.size = Pt(8.5)
            phead.bold = True
            phead.font.color.rgb = CLR_MUTED
            for prec in precs[:2]:
                rel = int(prec.get("relevanceScore", 0) * 100)
                pr2 = pp2.add_run(
                    f"{prec.get('title','')} ({prec.get('court','')}, {prec.get('year','')}) [{rel}% match]  "
                )
                pr2.font.size = Pt(8.5)
                pr2.font.color.rgb = CLR_PRIMARY

        # ── AI Rewrite box ─────────────────────────────────────────────────
        rewrite = clause.get("rewrite")
        if rewrite:
            rwtbl = doc.add_table(rows=1, cols=1)
            _no_border_tbl(rwtbl)
            rwc = rwtbl.rows[0].cells[0]
            _shade_cell(rwc, "F0FDF4")
            _cell_pad(rwc, 120, 120, 180, 180)

            rwp = rwc.paragraphs[0]
            rwp.paragraph_format.space_after = Pt(4)
            rwh = rwp.add_run("✦  AI RECOMMENDED SAFER VERSION")
            rwh.bold = True
            rwh.font.size = Pt(8.5)
            rwh.font.color.rgb = RISK_CLR["safe"]

            rwb = rwc.add_paragraph()
            rwb.paragraph_format.space_before = Pt(4)
            rw_run = rwb.add_run(str(rewrite))
            rw_run.font.size = Pt(9.5)
            rw_run.font.color.rgb = RGBColor(0x06, 0x4E, 0x3B)

        # Divider between clauses
        if idx < len(sorted_clauses) - 1:
            _hr(doc, "F3F4F6", before=Pt(10), after=Pt(10), size="4")


# ── Main Entry Point ───────────────────────────────────────────────────────────

def generate_report(data: dict) -> bytes:
    """
    Generate a professional legal audit DOCX report.

    Args:
        data: AnalysisResult dict with clauses, conflicts, summary, etc.

    Returns:
        DOCX file as bytes.
    """
    if not data or not isinstance(data, dict):
        raise ValueError("Invalid analysis data")
    if "clauses" not in data:
        raise ValueError("Missing clauses in analysis data")

    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Global typography ───────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = CLR_DARK
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    # ── Build sections ──────────────────────────────────────────────────────
    _build_cover(doc, data)
    doc.add_page_break()

    _build_summary(doc, data)

    conflicts = data.get("conflicts") or []
    if conflicts:
        _build_conflicts(doc, data)

    _build_clauses(doc, data)

    _add_footer(doc, str(data.get("documentId", "N/A")))

    # ── Serialize ───────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
